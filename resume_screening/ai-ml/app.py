from flask import Flask, request, jsonify
from resume_parser import extract_skills
from skill_matcher import match
import pdfplumber
import os

app = Flask(__name__)

@app.route("/analyze", methods=["POST"])
def analyze():
    try:
        text = ""
        job_skills = []

        # ✅ Check if file is uploaded as multipart form-data
        if 'file' in request.files:
            file = request.files['file']
            filename = file.filename
            
            import io
            file_bytes = file.read()
            file_stream = io.BytesIO(file_bytes)
            
            print("AI RECEIVED UPLOADED FILE:", filename)

            if filename.lower().endswith(".pdf"):
                with pdfplumber.open(file_stream) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + " "
            elif filename.lower().endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(file_stream)
                    for p in doc.paragraphs:
                        text += p.text + " "
                except Exception as e:
                    print('DOCX parse error:', e)
                    return jsonify({"error": "Failed to parse .docx file"}), 400
            elif filename.lower().endswith('.doc'):
                return jsonify({"error": "Unsupported file type: .doc (use .pdf or .docx)"}), 400
            else:
                text = file_bytes.decode('utf-8', errors='ignore')

            # Job skills might be passed as a form field
            job_skills_raw = request.form.get('jobSkills')
            if job_skills_raw:
                import json
                try:
                    job_skills = json.loads(job_skills_raw)
                except Exception:
                    job_skills = []
        else:
            # ✅ Fallback to JSON payload with filePath (for local/backward compatibility)
            data = request.get_json()
            if not data or "filePath" not in data:
                return jsonify({"error": "No file uploaded and no filePath provided"}), 400

            path = data["filePath"]
            print("FILE PATH RECEIVED:", path)

            if not os.path.exists(path):
                return jsonify({"error": f"File not found: {path}"}), 400

            if path.lower().endswith(".pdf"):
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + " "
            elif path.lower().endswith('.docx'):
                try:
                    from docx import Document
                    doc = Document(path)
                    for p in doc.paragraphs:
                        text += p.text + " "
                except Exception as e:
                    print('DOCX parse error:', e)
                    return jsonify({"error": "Failed to parse .docx file"}), 400
            elif path.lower().endswith('.doc'):
                return jsonify({"error": "Unsupported file type: .doc (use .pdf or .docx)"}), 400
            else:
                with open(path, "r", errors="ignore") as f:
                    text = f.read()

            job_skills = data.get('jobSkills') if (isinstance(data, dict) and data.get('jobSkills') is not None) else []

        if not text.strip():
            return jsonify({"error": "Empty resume content"}), 400

        # ✅ Skill extraction
        resume_skills = extract_skills(text)

        score, matched, missing = match(resume_skills, job_skills)

        # Debugging
        print("RESUME SKILLS:", resume_skills)
        print("MATCHED:", matched, "MISSING:", missing, "SCORE:", score)

        return jsonify({
            "name": "Candidate",
            "skills": resume_skills,
            "matchedSkills": matched,
            "jobSkills": job_skills,
            "experience": 2,
            "score": score,
            "missingSkills": missing
        })

    except Exception as e:
        print("AI ERROR:", e)
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
